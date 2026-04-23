"""
Build a polished Word document from the backend checklist.
Uses python-docx with a Kairox brand palette.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────── Palette ───────────────
NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x25, 0x63, 0xEB)
TEAL = RGBColor(0x10, 0xA3, 0x7F)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
DANGER = RGBColor(0xEF, 0x44, 0x44)
SLATE = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = "F1F5F9"
CODE_BG = "0B1220"
CODE_FG = RGBColor(0xE2, 0xE8, 0xF0)
ACCENT_BG = "EFF6FF"
SUCCESS_BG = "ECFDF5"
WARN_BG = "FFFBEB"
DANGER_BG = "FEF2F2"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

# ─────────────── Helpers ───────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, color="E2E8F0", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = color
    # Accent bar via a thin table underneath
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(3)
    set_cell_bg(cell, "2563EB")
    cell.text = ""
    for para in cell.paragraphs:
        para.paragraph_format.line_spacing = Pt(2)
        for r in para.runs:
            r.font.size = Pt(1)
    tbl.rows[0].height = Pt(3)
    return p

def add_h2(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = color
    return p

def add_h3(text, color=SLATE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = color
    return p

def add_para(text, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(10.5)
    return p

def add_callout(title, body, bg=ACCENT_BG, accent=BLUE, icon="ℹ"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_bg(cell, bg)
    set_cell_border(cell, color="DBEAFE")
    # Clear default paragraph
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"{icon}  {title}")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = accent
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = SLATE
    p2.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()  # spacer

def add_code_block(code):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_bg(cell, CODE_BG)
    set_cell_border(cell, color="1F2937")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.split("\n")):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
    doc.add_paragraph()

def add_table(headers, rows, col_widths=None, header_bg="0F172A",
              header_fg=RGBColor(0xFF, 0xFF, 0xFF), zebra_bg="F8FAFC",
              first_col_bold=False):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        if col_widths:
            c.width = col_widths[i]
        set_cell_bg(c, header_bg)
        set_cell_border(c, color="0F172A", sz="6")
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = header_fg
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Rows
    for ri, row in enumerate(rows):
        r = tbl.rows[ri + 1]
        for i, v in enumerate(row):
            c = r.cells[i]
            if col_widths:
                c.width = col_widths[i]
            if ri % 2 == 1:
                set_cell_bg(c, zebra_bg)
            set_cell_border(c, color="E2E8F0")
            c.text = ""
            p = c.paragraphs[0]
            run = p.add_run(str(v))
            run.font.size = Pt(9.5)
            if first_col_bold and i == 0:
                run.font.bold = True
                run.font.color.rgb = NAVY
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)

def page_break():
    doc.add_page_break()

# ─────────────── COVER PAGE ───────────────
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.autofit = False
cover_cell = cover_tbl.rows[0].cells[0]
cover_cell.width = Cm(17)
set_cell_bg(cover_cell, "0F172A")
cover_cell.text = ""
# Brand tag
p = cover_cell.paragraphs[0]
p.paragraph_format.space_before = Pt(60)
r = p.add_run("KAIROX  •  AI OPEX")
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
# Title
p = cover_cell.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run("Backend Build\nChecklist")
r.font.size = Pt(44)
r.font.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
# Subtitle
p = cover_cell.add_paragraph()
p.paragraph_format.space_before = Pt(16)
r = p.add_run("v3.0 — Delta spec for the AWS backend to make the mocked v3.0 frontend real.")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
# Pills
p = cover_cell.add_paragraph()
p.paragraph_format.space_before = Pt(36)
for label, color in [("17 sections", "2563EB"), ("~60 endpoints", "10A37F"), ("10 new tables", "F59E0B"), ("4 roles", "DB2777")]:
    r = p.add_run(f"  {label}  ")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Just colour the text distinctly; docx doesn't support inline pill easily
    r.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD) if color == "2563EB" else (
        RGBColor(0x6E, 0xE7, 0xB7) if color == "10A37F" else (
        RGBColor(0xFC, 0xD3, 0x4D) if color == "F59E0B" else RGBColor(0xF9, 0xA8, 0xD4)
    ))
    r2 = p.add_run("   ")
# Footer meta
p = cover_cell.add_paragraph()
p.paragraph_format.space_before = Pt(120)
r = p.add_run("Version 3.0  ·  April 2026\nSource of truth for all [BACKEND-GAP] warnings in the frontend.")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
p = cover_cell.add_paragraph()
r = p.add_run(" ")
r.font.size = Pt(20)
# Pad bottom
cover_cell.paragraphs[-1].paragraph_format.space_after = Pt(60)

page_break()

# ─────────────── EXECUTIVE SUMMARY ───────────────
add_h1("Executive Summary")

add_para(
    "The Kairox AI OpEx v3.0 React Native frontend has shipped all 17 feature sections "
    "across 4 user roles (Problem Solver, Supervisor, Managing Director, Customer's MD). "
    "Every data flow is currently mock-driven. This document is the exact delta your existing "
    "AWS backend needs to absorb so the frontend can be flipped from mock calls to real HTTPS "
    "endpoints, one feature at a time, without any further frontend rewrites."
)

add_callout(
    "The mock file IS the spec.",
    "Every /app/frontend/src/services/mocks/*.js file is ≤250 lines and exactly mirrors the "
    "required response shape. When a real endpoint ships, swap one import in services/api.js "
    "and the mock branch drops out — no other frontend change needed.",
    bg=ACCENT_BG, accent=BLUE, icon="💡"
)

add_h2("At a glance")
add_table(
    ["Area", "What ships", "Complexity"],
    [
        ["Roles & auth", "Add customer_md enum, junction table, JWT payload extend", "🟢 Low"],
        ["Directory APIs", "5 read endpoints — supervisors, CMDs, team, MD profile", "🟢 Low"],
        ["Personal chat (§7, §8)", "2 tables, 7 endpoints, role-pair access matrix, WS", "🟡 Medium"],
        ["Group chat + chatbot (§14, §2)", "3 tables, 6 endpoints, NLP intents, monthly AI summary", "🟡 Medium"],
        ["Issue lifecycle + escalation (§3)", "Status enum, escalation table + endpoint, email/chat triggers", "🟢 Low"],
        ["Alerts: Voice + WhatsApp (§5, §6)", "asyncio.gather dual-channel, Twilio webhooks, missed-call fallback", "🔴 High"],
        ["Photo timeline (§16)", "Add DURING enum, AI moderation flag, 3 endpoints", "🟢 Low"],
        ["CMD dashboard (§9)", "Single aggregator endpoint with site scoping", "🟢 Low"],
        ["Budget state machine (§11)", "3 tables, 3-tier classifier, 10 endpoints, audit log, hourly beat", "🟡 Medium"],
        ["Site diary + monthly report (§15)", "Table, 5 endpoints, LLM highlights, PDF export", "🟡 Medium"],
        ["MD admin (§12)", "4 endpoints — sites, users, CMD assignments", "🟢 Low"],
        ["Google Sheets sync (§13)", "OAuth, webhook, cell-diff worker, incremental sync", "🔴 High"],
        ["Project timeline (§17)", "2 tables, cascade algorithm, 4 endpoints, role scoping", "🟡 Medium"],
    ],
    col_widths=[Cm(5.5), Cm(9), Cm(2.5)],
    first_col_bold=True,
)

add_callout(
    "Recommended sequence",
    "Week 1: Roles + MD admin (unlocks user/site creation).  "
    "Week 2: Alerts + escalations.  "
    "Week 3: Budget + site diary.  "
    "Week 4: Personal + group chat + chatbot.  "
    "Week 5: Google Sheets.  "
    "Week 6: Gantt cascade + polish.",
    bg=SUCCESS_BG, accent=TEAL, icon="📅"
)

page_break()

# ─────────────── SECTION 1 — GLOBAL ───────────────
add_h1("1. Global Changes")

add_h2("1.1 Role model")
add_table(
    ["Change", "Detail"],
    [
        ["Add enum value", "ALTER TYPE user_role ADD VALUE 'customer_md';"],
        ["Alias", "manager → managing_director (keep manager for backward compat)"],
        ["New junction table", "customer_md_sites (user_id, site_id, assigned_at, assigned_by)"],
    ],
    col_widths=[Cm(5), Cm(12)],
    first_col_bold=True,
)

add_h2("1.2 Auth")
add_bullet("POST /api/v1/auth/login — extend to accept & return role: customer_md.")
add_bullet("Passwords: bcrypt (cost 12). Do not roll your own.")
add_bullet("JWT payload must include role, user_id, assigned_site_ids for fast FE scoping.")
add_bullet("Rate-limit: 5 attempts / 10 min per username.")
add_bullet("No signup / email-invite. MD creates users via §12 admin and shares credentials manually.")

add_h2("1.3 Shared response shape")
add_bullet("Lists: { items: [...], next_cursor: string | null, total?: number }")
add_bullet("Mutations: return the full updated resource, not just an id.")
add_bullet("Timestamps: ISO-8601 UTC.")
add_bullet("Currency: store in paise (integer). Frontend formats ₹ with Indian grouping.")

add_h2("1.4 Authorisation middleware")
add_para(
    "Build ONE reusable decorator: @require_role(['manager']) / @require_site_access(site_id). "
    "It reads JWT, applies the site-access matrix below, returns 403 {error: \"site-not-in-scope\"} on deny.",
    italic=True, color=MUTED,
)
add_table(
    ["Role", "Site access rule"],
    [
        ["manager", "Always allowed"],
        ["supervisor", "site_id IN supervisor_sites[me]"],
        ["customer_md", "site_id IN customer_md_sites[me]"],
        ["problem_solver", "Assigned to any issue in that site"],
    ],
    col_widths=[Cm(4), Cm(13)],
    first_col_bold=True,
)

page_break()

# ─────────────── SECTION 2 — DIRECTORY ───────────────
add_h1("2. Directory / People APIs")

add_table(
    ["Method", "Path", "Caller", "Purpose"],
    [
        ["GET", "/api/v1/supervisors", "MD", "All supervisors + stats (active, closed, sites)"],
        ["GET", "/api/v1/supervisors/:id", "MD / Sup (self)", "Profile + assigned sites + issue counts"],
        ["GET", "/api/v1/customer-mds", "MD", "All CMDs + company + assigned site names"],
        ["GET", "/api/v1/customer-mds/:id", "MD / CMD (self)", "Customer MD profile"],
        ["GET", "/api/v1/users", "MD only", "Flat list for Team screen"],
        ["GET", "/api/v1/me/md", "Sup / CMD", "Resolve the 'company MD' for the MD-card screen"],
    ],
    col_widths=[Cm(1.8), Cm(5), Cm(3), Cm(7.2)],
)

add_callout(
    "Shape gotchas",
    "supervisors response must include active_issues_count, closed_issues_count, sites:[{id,name}].  "
    "customer-mds response must include company and assigned_sites derived from the new junction table.",
    bg=ACCENT_BG, accent=BLUE, icon="⚠",
)

page_break()

# ─────────────── SECTION 3 — PERSONAL CHAT ───────────────
add_h1("3. Personal Chat (§7 Sup↔MD, §8 MD↔CMD, §10 access rules)")

add_h2("3.1 Schema")
add_code_block("""personal_chat_threads
  id (uuid)
  user_a_id (fk users.id)           -- smaller id always goes in user_a
  user_b_id (fk users.id)
  last_message_at
  created_at
  UNIQUE(user_a_id, user_b_id)

personal_chat_messages
  id (uuid)
  thread_id (fk)
  sender_id (fk users.id)
  body (text)
  type (enum: 'text','budget_card','site_diary','summary','pin_decision')
  payload (jsonb)                   -- e.g. budget_card:{amount,reason,site_id,status}
  is_pinned_decision (bool)         -- MD only
  reply_to_message_id (uuid)
  created_at, edited_at, deleted_at""")

add_h2("3.2 Endpoints")
add_table(
    ["Method", "Path", "Caller", "Behaviour"],
    [
        ["GET",   "/api/v1/personal-chats", "Any", "List conversations for current user"],
        ["GET",   "/api/v1/personal-chats/:id/messages", "Member", "Paged, newest→oldest"],
        ["POST",  "/api/v1/personal-chats", "Any", "Body: {peer_user_id}. Upsert thread"],
        ["POST",  "/api/v1/personal-chats/:id/messages", "Member", "Body: {body,type,payload?,reply_to_message_id?}"],
        ["POST",  "/api/v1/personal-chats/:id/messages/:msg/pin", "MD only", "Set is_pinned_decision=true, broadcast on WS"],
        ["POST",  "/api/v1/personal-chats/:id/messages/:msg/unpin", "MD only", "Opposite"],
        ["DELETE","/api/v1/personal-chats/:id/messages/:msg", "Sender", "Soft-delete"],
    ],
    col_widths=[Cm(2), Cm(6.5), Cm(2.5), Cm(6)],
)

add_h2("3.3 Role-pair access matrix (§10)")
add_para("Enforced in the thread-creation endpoint; return 403 {error:\"chat-not-allowed-between-roles\"} on deny.", italic=True, color=MUTED)
add_table(
    ["From → To", "Allowed?"],
    [
        ["supervisor → manager", "✅"],
        ["manager → supervisor", "✅"],
        ["manager → customer_md", "✅"],
        ["customer_md → manager", "✅"],
        ["supervisor → customer_md", "❌"],
        ["customer_md → supervisor", "❌"],
        ["problem_solver → *", "❌ (PS has no personal chat in v3.0)"],
        ["* → problem_solver", "❌"],
    ],
    col_widths=[Cm(6), Cm(11)],
    first_col_bold=True,
)

add_h2("3.4 WebSocket")
add_bullet("/ws/personal-chats/:thread_id — bidirectional JSON frames.")
add_bullet("On new message: {event:\"message\", message:{...}} to both members.")
add_bullet("On pin: {event:\"pin\", message_id, pinned:true}.")
add_bullet("Fallback: frontend polls every 3 s if WS drops. Keep REST list idempotent.")

page_break()

# ─────────────── SECTION 4 — GROUP CHAT + CHATBOT ───────────────
add_h1("4. Group Chat + Chatbot + AI Summaries")

add_h2("4.1 Operations Team group (§14)")
add_code_block("""group_chats(id, name='operations_team', created_by, created_at)
group_chat_members(group_id, user_id, role_in_group, joined_at)
group_chat_messages(... same shape as personal_chat_messages ... + group_id)
group_chat_summaries(id, group_id, month 'YYYY-MM', summary_text, generated_at)""")

add_table(
    ["Method", "Path", "Purpose"],
    [
        ["GET",  "/api/v1/group-chats/ops", "Return Ops Team group (auto-create with MD + all supervisors)"],
        ["GET",  "/api/v1/group-chats/:id/messages", "Paged messages"],
        ["POST", "/api/v1/group-chats/:id/messages", "Post (type: text / budget_card / ai_summary)"],
        ["POST", "/api/v1/group-chats/:id/messages/:msg/pin", "MD-only pin"],
        ["GET",  "/api/v1/group-chats/:id/summaries", "Last 6 monthly AI summaries"],
        ["POST", "/api/v1/group-chats/:id/summaries/regenerate", "MD queues a re-run"],
    ],
    col_widths=[Cm(1.8), Cm(7.5), Cm(7.7)],
)
add_para("WebSocket: /ws/group-chats/:id (same protocol as personal).", italic=True, color=MUTED)

add_h2("4.2 Role-personalised chatbot (§2)")
add_para("Extend your existing POST /api/v1/chat endpoint:")
add_code_block("""Request:  { session_id, message, role, user_id }
Response: { reply, data: { intent, entities, quick_actions[], channels? } }""")

add_para("Add 3 NLP intents (wire into your existing intent router):")
add_table(
    ["Intent", "Example message", "Handler"],
    [
        ["budget_request", "Raise 2 lakh for new sensors at Zone C",
         "Classify budget → insert budget_requests → return budget_card payload"],
        ["budget_report", "Show top 5 escalated budgets",
         "Aggregate from budget_requests → return markdown summary"],
        ["site_diary_entry", "Log today: replaced pump at Ambattur, 2 h rain",
         "Insert site_diary_entries → return confirmation card"],
    ],
    col_widths=[Cm(3.5), Cm(6.5), Cm(7)],
    first_col_bold=True,
)

add_h2("4.3 Monthly summary Celery beat")
add_bullet("Schedule: 1st of every month, 02:00 IST.")
add_bullet("For each group chat + each personal thread with ≥10 messages, call LLM with last month's messages.")
add_bullet("Insert into group_chat_summaries AND post a type=ai_summary message in the chat.")

page_break()

# ─────────────── SECTION 5 — ISSUE LIFECYCLE ───────────────
add_h1("5. Issue Lifecycle + Escalation (§3)")

add_h2("5.1 Status enum → Kairox label")
add_table(
    ["Backend enum", "Kairox label (FE)", "Colour"],
    [
        ["OPEN / ASSIGNED / AUTO_ASSIGNED / REASSIGNED", "Active", "Blue"],
        ["IN_PROGRESS", "In Progress", "Amber"],
        ["COMPLETED", "Fixed", "Green"],
        ["ESCALATED", "Escalated", "Red"],
        ["REOPENED", "Not Fixed", "Red"],
    ],
    col_widths=[Cm(8), Cm(5), Cm(4)],
    first_col_bold=True,
)
add_para("No DB change needed; just ensure exact enum strings. Frontend owns the label map.", italic=True, color=MUTED)

add_h2("5.2 Escalation report")
add_code_block("""issue_escalations
  id, issue_id (fk),
  raised_by_supervisor_id,
  reason (text),
  root_cause (text), proposed_action (text),
  copy_customer_md (bool),
  email_sent_at, chat_notification_sent_at,
  created_at""")

add_table(
    ["Method", "Path", "Caller", "Behaviour"],
    [
        ["POST", "/api/v1/issues/:id/escalate", "Supervisor",
         "Flip issue to ESCALATED, insert row, email MD (+CMD if copy), post chat notification"],
        ["GET",  "/api/v1/issues/:id/escalations", "MD / Sup / CMD (access)",
         "Return all escalation records for the issue"],
    ],
    col_widths=[Cm(1.8), Cm(6.2), Cm(3), Cm(6)],
)

page_break()

# ─────────────── SECTION 6 — ALERTS ───────────────
add_h1("6. Dual-Channel Alerts (§5 + §6)")

add_h2("6.1 Flow when a Supervisor creates an issue")
add_para("Fire voice + WhatsApp in parallel using asyncio.gather:")
add_code_block("""async def send_dual_channel_alert(issue, recipient_phone):
    voice, whatsapp = await asyncio.gather(
        twilio.place_alert_call(issue, recipient_phone),
        whatsapp.send_template(recipient_phone, TEMPLATE_NEW_ISSUE, vars),
        return_exceptions=True,
    )
    persist_issue_alert(issue.id, voice, whatsapp)
    response.data.channels = {"voice_call": voice.status,
                              "whatsapp":   whatsapp.status,
                              "all_ok":     voice.ok and whatsapp.ok}""")

add_h2("6.2 Persistence")
add_code_block("""issue_alerts
  id, issue_id (fk),
  voice_call_status (delivered|missed|failed|pending),
  voice_call_sid, voice_call_duration_sec,
  whatsapp_status (delivered|failed|pending),
  whatsapp_message_sid,
  all_ok (bool), missed_call (bool),
  created_at""")

add_h2("6.3 Missed-call fallback (§6)")
add_bullet("Twilio status webhook: /api/v1/webhooks/twilio/call-status (signature-verified, no JWT).")
add_bullet("On no-answer / busy → set missed_call=true, voice_call_status='missed'.")
add_bullet("Fire a DIFFERENT WhatsApp template to PS + Supervisor: 'Call missed — please check app. Issue #{id} {title}'.")
add_bullet("Broadcast {event:\"alert-updated\"} on /ws/issues/:id so the banner flips amber in real time.")

add_h2("6.4 Endpoints")
add_table(
    ["Method", "Path", "Purpose"],
    [
        ["GET",  "/api/v1/issues/:id/alert", "Return issue_alerts row for AlertStatusBanner"],
        ["POST", "/api/v1/issues/:id/alert/whatsapp/resend", "Re-fire WhatsApp (Sup or MD)"],
        ["POST", "/api/v1/webhooks/twilio/call-status", "Twilio status webhook"],
    ],
    col_widths=[Cm(1.8), Cm(8.5), Cm(6.7)],
)

page_break()

# ─────────────── SECTION 7 — PHOTOS ───────────────
add_h1("7. Photo Progress Timeline (§16)")

add_h2("7.1 Enum extension")
add_code_block("ALTER TYPE image_phase ADD VALUE 'DURING';")

add_h2("7.2 Endpoints")
add_table(
    ["Method", "Path", "Access"],
    [
        ["GET",    "/api/v1/issues/:id/photos?group_by=phase", "Any with access. Returns {before:[], during:[], after:[]}"],
        ["POST",   "/api/v1/issues/:id/photos", "Supervisor=BEFORE, PS=DURING/AFTER"],
        ["DELETE", "/api/v1/issues/:id/photos/:photoId", "Uploader or MD"],
    ],
    col_widths=[Cm(1.8), Cm(7.5), Cm(7.7)],
)

add_h2("7.3 AI validation on AFTER photos")
add_bullet("Run uploaded AFTER photo through your image-moderation / sanity-check pipeline.")
add_bullet("If flagged, set issue_photos.ai_flag=true and call NotificationService.send_photo_flag_alert.")
add_bullet("Post a text message into the Sup↔PS chat AND add a banner in issue detail.")

page_break()

# ─────────────── SECTION 8 — CMD DASHBOARD ───────────────
add_h1("8. Customer MD Dashboard (§9)")

add_para("Single aggregator endpoint the frontend calls on CMD login.")
add_para("GET /api/v1/dashboard/customer-md (CMD only)", bold=True, color=NAVY)

add_code_block("""{
  "stats": {
    "pending_issues": 7,
    "resolved_issues": 15,
    "total_sites": 3,
    "escalated_to_me": 2
  },
  "sites": [
    {"id":1,"name":"…","location":"…","issues_total":12,"issues_open":4}, …
  ],
  "escalated_budgets": [
    {"id":…,"amount":…,"site_name":"…","reason":"…"}, …
  ],
  "recent_issues": [ … limit 10 …]
}""")

add_callout(
    "Scoping rule (applies everywhere CMD sees data)",
    "site_id IN (SELECT site_id FROM customer_md_sites WHERE user_id = :me)",
    bg=WARN_BG, accent=AMBER, icon="🔒"
)

page_break()

# ─────────────── SECTION 9 — BUDGET ───────────────
add_h1("9. Budget — Multi-level Approval (§11)")

add_h2("9.1 Schema")
add_code_block("""site_budgets
  site_id (pk), monthly_ceiling (int, paise), updated_at

budget_requests
  id (uuid), title, reason, amount (int, paise),
  site_id, raised_by_user_id, raised_at,
  status (pending_md|approved|rejected|escalated_customer_md|
          cmd_approved|cmd_rejected|auto_approved),
  decided_by_user_id, decided_at, decision_note

budget_audit_log
  id, budget_request_id, actor_user_id,
  action (created|auto_approved|md_approved|md_rejected|
          escalated|cmd_approved|cmd_rejected),
  note, created_at

budget_thresholds  (admin-editable, per-company)
  company_id (pk), auto_limit, md_limit""")

add_h2("9.2 Classification logic")
add_code_block("""if amount <= AUTO_LIMIT (default ₹50,000):
    status = 'auto_approved'
elif amount <= MD_LIMIT   (default ₹2,00,000):
    status = 'pending_md'
else:
    status = 'pending_md'   # MD can still reject outright;
                            # MD's "escalate" action → 'escalated_customer_md'""")

add_h2("9.3 State machine")
add_code_block("""               ┌─ auto_approved          (amount ≤ AUTO_LIMIT)
pending_md ────┼─ approved               (md_accept)
               ├─ rejected               (md_reject)
               └─ escalated_customer_md  (md_escalate, only if amount > MD_LIMIT)
                         ├─ cmd_approved (cmd_esc_approve)
                         └─ cmd_rejected (cmd_esc_reject)""")

add_h2("9.4 Endpoints")
add_table(
    ["Method", "Path", "Caller", "Purpose"],
    [
        ["POST", "/api/v1/budget/requests", "Sup", "Create request (classifies, auto-approves small)"],
        ["GET",  "/api/v1/budget/requests?scope=<role>", "Sup/MD/CMD", "Role-scoped list"],
        ["GET",  "/api/v1/budget/totals?scope=<role>", "same", "{count, pending, approvedSum, rejectedCount}"],
        ["POST", "/api/v1/budget-requests/:id/accept", "MD", "→ approved + audit"],
        ["POST", "/api/v1/budget-requests/:id/reject", "MD", "→ rejected + audit"],
        ["POST", "/api/v1/budget-requests/:id/escalate", "MD", "→ escalated_customer_md (amount > MD_LIMIT)"],
        ["POST", "/api/v1/budget-requests/:id/esc-approve", "CMD", "→ cmd_approved"],
        ["POST", "/api/v1/budget-requests/:id/esc-reject", "CMD", "→ cmd_rejected"],
        ["GET",  "/api/v1/budgets/sites/:id", "Sup/MD/CMD (access)", "MTD approved vs ceiling"],
        ["GET",  "/api/v1/budgets/sites/:id/history", "same", "Monthly series for charts"],
        ["GET",  "/api/v1/budget/burn-rate", "Sup/MD", "Per-site {site_id, spent, ceiling, ratio}"],
        ["GET",  "/api/v1/budgets/threshold-alerts", "MD", "Sites at ≥90% of ceiling"],
    ],
    col_widths=[Cm(1.6), Cm(6.2), Cm(2.5), Cm(6.7)],
)

add_h2("9.5 Threshold-alert Celery beat")
add_bullet("Hourly: for each site, compute this month's burn.")
add_bullet("If ≥90% AND no alert fired in last 24 h → insert notification + post message in Ops Team group.")

page_break()

# ─────────────── SECTION 10 — SITE DIARY + MONTHLY REPORT ───────────────
add_h1("10. Site Diary (§15) + Monthly Report")

add_h2("10.1 Schema")
add_code_block("""site_diary_entries
  id (uuid), site_id (fk),
  author_id (fk users.id),          -- Problem Solver
  entry_date (date),                -- ideally one per (site, author, date)
  work_done (text),
  issues_noted (text),
  weather (text),                   -- optional, auto-fill via weather API
  safety_incidents (text default 'None'),
  photos (jsonb),                   -- array of photo URIs
  created_at""")

add_h2("10.2 Endpoints")
add_table(
    ["Method", "Path", "Caller", "Purpose"],
    [
        ["POST", "/api/v1/site-diary", "PS", "Create entry"],
        ["GET",  "/api/v1/site-diary", "Sup/MD/CMD/PS", "List with filters (site_id, author_id, cmd_sites)"],
        ["GET",  "/api/v1/site-diary/:id", "any (access)", "Fetch one"],
        ["GET",  "/api/v1/site-diary/monthly-report", "Sup/MD/CMD", "The aggregator — see 10.3"],
        ["GET",  "/api/v1/site-diary/export/pdf", "MD / CMD", "Server-side PDF (WeasyPrint / ReportLab)"],
    ],
    col_widths=[Cm(1.8), Cm(6.5), Cm(2.5), Cm(6.2)],
)

add_h2("10.3 Monthly Report response contract")
add_para("Frontend currently builds this in monthlyReportMockService.js. Backend must return the same shape, role-scoped:", italic=True, color=MUTED)
add_code_block("""{
  "month": "April 2026",
  "scope_label": "Vikram Singh · All 5 sites · 3 supervisors",
  "role": "manager",
  "hero": { "label": "Closure rate", "value": "30%" },
  "kpis": [
    {"key":"raised","label":"Raised","value":27,"delta":"+23%",
     "trend":"up","invertSentiment":false}, …
  ],
  "per_site": [
    {"site_id":1,"name":"…","issues_total":6,"issues_closed":2,
     "issues_open":4,"on_time_rate":33,"spent":62000,
     "ceiling":1000000,"burn_ratio":0.062}, …
  ],
  "top_issues": [ {"id":7,"title":"…","status":"ESCALATED",
                   "priority":"high","site_name":"…",
                   "impact_score":86}, … ],
  "budget": {"approved_count":4,"approved_sum":380000,
             "rejected_count":1,"escalated_count":2},
  "standout": {"label":"Top supervisor","name":"Priya Sharma",
               "detail":"4 issues closed"},
  "safety_incidents": 0,
  "highlights": "Company closure rate is 30% — +60% vs last month. …"
}""")

add_h3("Logic")
add_bullet("Scope by role: supervisor→supervisor_sites, customer_md→customer_md_sites, manager→all.")
add_bullet("Deltas: query same metrics for previous month; compute (cur - prev) / prev × 100.")
add_bullet("hero.value differs per role — see the mock for exact mapping.")
add_bullet("highlights: LLM-generated. Pass KPIs + standout into a prompt. Cache per (scope, month) for 1 h.")
add_bullet("weather on diary entries: call OpenWeatherMap once per (site, date) using site.coordinates; store result.")

page_break()

# ─────────────── SECTION 11 — MD ADMIN ───────────────
add_h1("11. MD Admin — Sites, Users, CMD Assignments (§12)")

add_table(
    ["Method", "Path", "Body", "Behaviour"],
    [
        ["POST", "/api/v1/sites", "{name, location, latitude, longitude, initial_budget}",
         "Insert sites row + site_budgets(site_id, monthly_ceiling=initial_budget)"],
        ["POST", "/api/v1/users", "{name, phone, email, role, password, company?, site_ids?}",
         "Validate role. username=last 10 digits of phone. Hash password. Seed customer_md_sites if CMD."],
        ["GET", "/api/v1/users", "—", "Flat list for the Team screen"],
        ["POST", "/api/v1/customer-mds/:id/sites", "{site_ids: [int]}",
         "REPLACE entire assignment — delete for user_id, insert new set"],
    ],
    col_widths=[Cm(1.6), Cm(5.8), Cm(4.5), Cm(5)],
)

add_h2("Validation")
add_bullet("Phone must match E.164-ish: +\\d{10,15}")
add_bullet("Email: RFC-5322 (optional field).")
add_bullet("Password min 8 chars.")

page_break()

# ─────────────── SECTION 12 — GOOGLE SHEETS ───────────────
add_h1("12. Google Sheets Live Sync (§13)")

add_callout(
    "Biggest pure-backend build",
    "Frontend has 3 small endpoints + an animated status pill. Everything else (OAuth, webhook, "
    "cell-diff worker, column mapping, validation, incremental sync, token refresh, backoff) is "
    "server-side. Allocate ~1 week, 1 backend engineer.",
    bg=WARN_BG, accent=AMBER, icon="🏗"
)

add_h2("12.1 Schema")
add_code_block("""sheet_integrations
  id, user_id (MD who connected), sheet_id, sheet_name, spreadsheet_url,
  oauth_refresh_token (ENCRYPTED), webhook_channel_id,
  connected_at, disconnected_at, last_error (text)

sheet_sync_state
  integration_id (pk)
  last_synced_at, last_full_sync_at,
  record_count, state (idle|syncing|synced|error)

sheet_cell_diff   -- ring buffer, keep last 10k
  id, integration_id, row_index, col_index,
  old_value, new_value, applied_at, record_affected""")

add_h2("12.2 Endpoints")
add_table(
    ["Method", "Path", "Caller", "Purpose"],
    [
        ["POST", "/api/v1/integrations/google-sheets/connect", "MD", "Returns OAuth URL"],
        ["GET",  "/api/v1/integrations/google-sheets/oauth-callback", "Google", "Completes OAuth, stores refresh token, registers webhook"],
        ["GET",  "/api/v1/integrations/google-sheets/status", "MD", "Full state object (pill consumes this)"],
        ["POST", "/api/v1/integrations/google-sheets/sync", "MD", "Manual trigger — enqueue full sync"],
        ["POST", "/api/v1/integrations/google-sheets/disconnect", "MD", "Revoke token, delete webhook, clear state"],
        ["POST", "/api/v1/webhooks/google-sheets/push", "Google", "Webhook target → enqueue cell-diff job"],
    ],
    col_widths=[Cm(1.6), Cm(8), Cm(2.2), Cm(5.2)],
)

add_h2("12.3 Core logic")
add_bullet("OAuth2 with scope https://www.googleapis.com/auth/spreadsheets.")
add_bullet("Register a Drive API 'changes.watch' webhook so Google POSTs you on any edit.")
add_bullet("Cell-diff worker (Celery): batchGet current state → diff vs last snapshot (S3/Redis) → map changed cells to DB fields via column-mapping config → apply with validation.")
add_bullet("Refuse to sync a row if required columns are empty or types mismatch.")
add_bullet("Incremental only — NEVER full-row overwrites.")
add_bullet("Resilience: refresh expired tokens; exponential backoff on Google 429s.")

add_h2("12.4 Required env vars")
add_code_block("""GOOGLE_OAUTH_CLIENT_ID=…
GOOGLE_OAUTH_CLIENT_SECRET=…
GOOGLE_SHEETS_WEBHOOK_URL=https://api.kairoxaitech.com/api/v1/webhooks/google-sheets/push
SHEETS_INTEGRATION_SECRET=…   # HMAC signing for webhook payloads""")

page_break()

# ─────────────── SECTION 13 — TIMELINE ───────────────
add_h1("13. Project Timeline — Gantt + Dependency Cascade (§17)")

add_h2("13.1 Schema")
add_code_block("""site_contracts
  site_id (pk, fk), start_date, end_date, updated_at

site_tasks
  id (uuid), site_id (fk),
  name (text),
  start_date, end_date,
  assigned_to_supervisor_id (fk),
  status (not_started|in_progress|completed|delayed),
  depends_on_task_id (uuid, fk self, nullable),
  created_at, updated_at""")

add_h2("13.2 Endpoints")
add_table(
    ["Method", "Path", "Caller", "Purpose"],
    [
        ["GET",    "/api/v1/sites/:id/timeline", "MD, Sup(assigned), CMD(assigned)",
         "{site_id, contract_start_date, contract_end_date, tasks:[…]}"],
        ["POST",   "/api/v1/sites/:id/tasks", "MD", "Create task"],
        ["PATCH",  "/api/v1/sites/:id/tasks/:taskId", "MD", "Update any field; triggers cascade"],
        ["DELETE", "/api/v1/sites/:id/tasks/:taskId", "MD", "Remove + reparent dependents (set FK null)"],
    ],
    col_widths=[Cm(1.8), Cm(6.5), Cm(4), Cm(4.7)],
)

add_h2("13.3 Cascade algorithm (port from projectTimelineMockService.updateTask)")
add_code_block("""BEGIN TRANSACTION
  delta = new_end_date - old_end_date
  if delta > 0:
      queue = [task.id]; seen = {}
      while queue not empty:
          cur = queue.pop()
          for child in tasks where depends_on_task_id == cur:
              if child.id in seen: continue     # cycle guard
              seen.add(child.id)
              child.start_date += delta
              child.end_date   += delta
              if child.status == 'not_started':
                  child.status = 'delayed'
              queue.push(child.id)
      new_contract_end = max(tasks.end_date)
      if new_contract_end > site_contracts.end_date:
          site_contracts.end_date = new_contract_end
COMMIT    -- rollback on any exception

Return { timeline, affected: [task_ids_shifted] }""")

add_callout(
    "Edge cases to test",
    "Moving end-date BACKWARDS (no cascade). Cycle in depends_on_task_id (guard must stop). "
    "Parallel dependents — both shift by the same delta. Final task extends beyond contract end — "
    "contract end updates atomically in the same txn.",
    bg=WARN_BG, accent=AMBER, icon="🧪"
)

page_break()

# ─────────────── SECTION 14 — WS + POLLING ───────────────
add_h1("14. WebSockets + Polling Fallback")

add_h2("14.1 Endpoints")
add_table(
    ["Path", "Purpose"],
    [
        ["/ws/personal-chats/:thread_id", "Messages, pins"],
        ["/ws/group-chats/:id", "Messages, pins, ai_summary"],
        ["/ws/issues/:id", "alert-updated, escalation events"],
        ["/ws/dashboard/:user_id", "threshold-alerts, sheet sync state"],
    ],
    col_widths=[Cm(6.5), Cm(10.5)],
    first_col_bold=True,
)

add_h2("14.2 Auth & fallback")
add_bullet("First frame on connect: {\"type\":\"auth\",\"token\":\"<JWT>\"}. Disconnect on invalid.")
add_bullet("If WS fails, frontend polls REST every 3 s. Honour If-Modified-Since so polling stays cheap.")

page_break()

# ─────────────── SECTION 15 — WORKERS ───────────────
add_h1("15. Background Workers (Celery / cron)")

add_table(
    ["Schedule", "Task", "Purpose"],
    [
        ["0 2 1 * * (monthly)", "generate_monthly_summaries", "LLM summary for every active chat"],
        ["0 * * * * (hourly)",  "check_budget_thresholds",    "Fire threshold-alerts at ≥90% site burn"],
        ["*/5 * * * *",         "google_sheets_resync_if_stale","Refresh if last_synced_at > 5 min ago"],
        ["On webhook",          "apply_sheet_cell_diff",       "Driven by Google changes.watch"],
        ["On issue create",     "send_dual_channel_alert",     "Voice + WhatsApp in parallel"],
        ["On escalation",       "send_escalation_email",       "+ post chat notification to MD"],
        ["Nightly",             "ai_flag_after_photos",        "Batch re-check flagged photos (rate-limit AI)"],
    ],
    col_widths=[Cm(3.5), Cm(5.5), Cm(8)],
)

page_break()

# ─────────────── SECTION 16 — INTEGRATIONS ───────────────
add_h1("16. External Integrations Summary")

add_table(
    ["Service", "Purpose", "Env vars"],
    [
        ["Twilio Voice API", "Auto-call Sup / PS on new issue",
         "TWILIO_SID, TWILIO_AUTH_TOKEN, TWILIO_FROM_NUMBER"],
        ["Twilio WhatsApp Business", "Templated alerts, missed-call fallback, CMD escalation notify",
         "TWILIO_WHATSAPP_FROM + 3 approved template SIDs"],
        ["Google Sheets v4 + Drive API", "Live sync (§13)",
         "GOOGLE_OAUTH_CLIENT_ID / SECRET, GOOGLE_SHEETS_WEBHOOK_URL, SHEETS_INTEGRATION_SECRET"],
        ["OpenWeatherMap", "weather field on site diary",
         "WEATHER_API_KEY"],
        ["LLM provider (Claude / GPT / Gemini)", "Monthly summaries, highlights, chatbot replies",
         "Already in place"],
        ["WeasyPrint / ReportLab", "Monthly-report PDF export",
         "—"],
        ["S3 / equivalent", "Photo uploads + sheet snapshot storage",
         "Already in place"],
        ["Sentry / log sink", "Webhook + cascade error tracking",
         "SENTRY_DSN"],
    ],
    col_widths=[Cm(4.2), Cm(6.8), Cm(6)],
    first_col_bold=True,
)

page_break()

# ─────────────── SECTION 17 — TEST MATRIX ───────────────
add_h1("17. Testing Matrix (minimum before promoting)")

add_table(
    ["#", "Test", "Type"],
    [
        ["1",  "Budget @ ₹30k → auto_approved; audit log has auto_approved action", "Integration"],
        ["2",  "Budget @ ₹1.5L → pending_md → md_accept → approved",                "Integration"],
        ["3",  "Budget @ ₹5L → pending_md → md_escalate → cmd_approve → cmd_approved", "Integration"],
        ["4",  "Sup raises issue → voice call + WhatsApp fire in parallel",          "Integration (mock Twilio)"],
        ["5",  "Twilio webhook no-answer → missed_call=true, fallback WhatsApp",     "Integration"],
        ["6",  "Shift task end-date +10d → dependents shift, contract end updates",  "Unit (cascade)"],
        ["7",  "Cascade with cycle in depends_on_task_id → loop terminates",         "Unit"],
        ["8",  "Google webhook POST → cell-diff applied to issues.title row N",      "Integration (mock Google)"],
        ["9",  "CMD login → dashboard shows only customer_md_sites",                 "Integration"],
        ["10", "MD adds CMD + 2 sites → CMD login shows exactly those 2 sites",      "E2E"],
        ["11", "MD pins a message → WS pin event → Sup sees it live",                "E2E"],
        ["12", "Monthly report for supervisor1 → scope_label matches name + sites",  "Integration"],
        ["13", "Delete task with dependents → FK cleared to null",                   "Unit"],
        ["14", "Rate-limit 6th login attempt → 429",                                 "Integration"],
    ],
    col_widths=[Cm(1), Cm(12.5), Cm(3.5)],
    first_col_bold=True,
)

page_break()

# ─────────────── SECTION 18 — ROLLOUT + CHEATSHEET ───────────────
add_h1("18. Rollout Plan & Cheatsheet")

add_h2("18.1 Suggested 6-week rollout")
add_table(
    ["Week", "Ship", "Unlocks"],
    [
        ["1", "§1 roles + §2 directory + §11 MD admin", "MD can onboard real users / sites"],
        ["2", "§5+§6 alerts + §3 escalations",          "Real voice + WhatsApp + escalation email"],
        ["3", "§9 budget + §15 site diary + monthly report", "End-to-end budget workflow + report"],
        ["4", "§7 + §8 + §14 chats + §2 chatbot intents", "Real-time chat + NLP intents"],
        ["5", "§13 Google Sheets sync",                  "Live sheet → DB propagation"],
        ["6", "§17 Gantt cascade + §16 photo AI flag + polish", "Full v3.0 parity"],
    ],
    col_widths=[Cm(1.5), Cm(8), Cm(7.5)],
    first_col_bold=True,
)

add_h2("18.2 Cheatsheet — find the exact frontend call-site")
add_para("Every mock file IS the spec. Point engineers at these greps:", italic=True, color=MUTED)
add_code_block("""# every backend-gap with its endpoint spec
grep -rn "BACKEND-GAP" /app/frontend/src

# every TODO(backend) comment above the call site
grep -rn "TODO(backend)" /app/frontend/src

# the canonical request/response shapes
ls /app/frontend/src/services/mocks/*.js""")

add_callout(
    "After each section is live",
    "Swap the corresponding mock import in src/services/mocks/* for a real axios call. "
    "The // TODO(backend): comments in every mock tell you exactly what URL + method. "
    "No further frontend change needed — contracts match 1-to-1.",
    bg=SUCCESS_BG, accent=TEAL, icon="✅"
)

# ─────────────── FOOTER ───────────────
add_divider()
p = doc.add_paragraph()
r = p.add_run("Kairox AI OpEx  •  Backend Build Checklist v3.0  •  April 2026")
r.font.size = Pt(8.5)
r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
out = "/app/Kairox_Backend_Build_Checklist_v3.docx"
doc.save(out)
print(f"Saved: {out}")
